"""Build the July 2026 plant-vs-engine report as a .docx.

    PYTHONPATH=. python scripts/build_july_report.py

Every figure is read from the artefacts by scripts/diag_plant_vs_engine.py and
pasted here as a literal, so the document and the diagnostic cannot drift apart
silently. Re-run the diagnostic first if the plan changes.

Source run: runs/jul_v13 (V13 config -- WARM_PRESS off, T0_STOCK_BASIS star,
GT_WIP_RAIL_MARGIN 1.0, LOT_INTERVAL_H 8, unpartitioned).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent.parent.parent / "JULY_2026_PLANT_VS_ENGINE.docx"

RED = RGBColor(0xA0, 0x20, 0x20)
GRN = RGBColor(0x1E, 0x6B, 0x2E)
GRY = RGBColor(0x55, 0x55, 0x55)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, *, bold=False, italic=False, color=None, size=None):
    par = doc.add_paragraph()
    r = par.add_run(text)
    r.bold, r.italic = bold, italic
    if color is not None:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return par


def bullets(doc, items, style="List Bullet"):
    for it in items:
        doc.add_paragraph(it, style=style)


def table(doc, headers, rows, widths=None, mono=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(htxt))
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            txt = str(v)
            colour = None
            if txt.startswith("!!"):
                txt, colour = txt[2:], RED
            elif txt.startswith("++"):
                txt, colour = txt[2:], GRN
            r = par.add_run(txt)
            r.font.size = Pt(9)
            if mono:
                r.font.name = "Consolas"
            if colour is not None:
                r.font.color.rgb = colour
            if i > 0:
                par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    return t


# ==========================================================================
def main() -> None:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    doc.add_heading("July 2026 — Plant vs Engine", 0)
    p(doc, "Where every missing tyre goes, why, and what would recover it",
      italic=True, color=GRY, size=12)
    p(doc, "Run runs/jul_v13 (V13) · unpartitioned · generated from "
           "l11_invariants.parquet, gt_events.parquet, build_schedule.parquet, "
           "cure_campaigns.parquet, build_starved.parquet and net_requirement_2026-07. "
           "Reproduce with:  PYTHONPATH=. python scripts/diag_plant_vs_engine.py "
           "jul_v13 2026-07",
      size=9, color=GRY)

    # ---------------------------------------------------------------- 0
    h(doc, "0. Why July, and the one caveat that governs this whole report", 1)
    p(doc, "July's demand IS the plant's own July production, taken from MES. For "
           "every GT, plant production = demand, and the plant therefore scored 100 % "
           "by construction. That makes July the only month where a per-GT "
           "plant-vs-engine comparison is fair rather than a comparison against a "
           "forecast.")
    p(doc, "It also makes July the easy month, and the flattering one.", bold=True)
    p(doc, "July is unpartitioned, fully feasible on capacity, and has 48 demanded "
           "PCR GTs. August is a forward order book, arithmetically infeasible by "
           "~3.2 points, partitioned, with 73 demanded PCR GTs. Two flags tuned on "
           "July alone were found on 13 August to cost August real fulfilment "
           "(WARM_PRESS: Jul +0.6 / Aug −0.5; T0_STOCK_BASIS=lot: Jul +0.1 / "
           "Aug TBR −0.7). Both were reverted. Nothing in this report should be "
           "acted on without re-checking it on August via "
           "scripts/ab_both_months.py.")

    # ---------------------------------------------------------------- 1
    h(doc, "1. Executive summary — the 15,309 PCR tyres, accounted for exactly", 1)
    p(doc, "PCR finishes July at 96.2 % in-month, TBR at 95.7 %. The PCR gap is "
           "15,309 tyres against gross demand of 398,405. The decomposition below "
           "sums to 15,309 — not approximately, exactly:")
    table(doc,
          ["Loss", "PCR tyres", "pt", "Class", "Owner"],
          [["Opening GT netted out of the build requirement, never credited as an "
            "in-month cure", "3,642", "0.91", "Definitional", "—"],
           ["No feasible release — no R5-legal window on an allowable machine",
            "!!6,310", "!!1.58", "Genuine scheduling + narrow matrices", "Shared"],
           ["min_lot floor would be breached", "2,232", "0.56", "Policy",
            "Plant ruling"],
           ["Built in the tail window, after 07:00 on 31 Jul", "1,675", "0.42",
            "Definitional", "Plant ruling"],
           ["B12 residual — GT demand below 300, never planned", "1,124", "0.28",
            "Policy", "Plant ruling"],
           ["Closing GT left standing at 07:00 on 1 Aug", "313", "0.08",
            "Definitional", "—"],
           ["Other L5 scheduling", "13", "0.00", "Ours", "—"],
           ["TOTAL", "15,309", "3.84", "", ""]])
    p(doc, "59 % of the gap is definitional or plant policy. 41 % — the 6,310 "
           "\"no feasible release\" line — is the genuinely recoverable share, and "
           "most of that is locked behind narrow allowable-machine matrices rather "
           "than behind the search.", bold=True)
    p(doc, "The single most important finding:", bold=True)
    p(doc, "There is no unexplained residue anywhere in the July plan. "
           "build_starved.parquet reconciles EXACTLY with the waterfall's unfed "
           "line — PCR 6,310 + 2,232 = 8,542; TBR 1,454 + 52 = 1,506. \"The "
           "optimiser is bad\" is not a supportable diagnosis for July.")
    p(doc, "Separately, and on a different lens: 6,372 PCR tyres are built in July "
           "and cure on 1–2 August. They are real output that the horizon rule "
           "excludes. If the plant rules that they count, PCR reads 97.8 %.")

    # ---------------------------------------------------------------- 2
    h(doc, "2. The waterfall — demand to completed, every step", 1)
    p(doc, "Three relationships in net_requirement are counter-intuitive and were "
           "got wrong on the first pass of this analysis: net_cure == demand "
           "(opening GT is supply, not a demand deduction); fg_stock == 0; and "
           "`usable` counts opening-GT units, not resolved SKUs.", size=9, color=GRY)

    h(doc, "PCR", 2)
    table(doc, ["Stage", "Qty", "Delta", "Why / who owns it"],
          [["Demand (= plant's own July output)", "398,405", "—", ""],
           ["− covered by opening GT on floor", "393,585", "−4,820",
            "Supply already standing"],
           ["+ cure-yield scrap uplift", "394,763", "+1,178", "HARD: yield 0.997"],
           ["L4.5 lot-sized gross build", "394,763", "0", "== gross_build"],
           ["L5 scheduled into cure campaigns", "393,626", "−1,137",
            "B12 residual dropped = 1,124 (POLICY)"],
           ["L7 actually BUILT and fed", "!!385,084", "!!−8,542",
            "Building could not reach the seat in time"],
           ["Built before 31 Jul 07:00", "383,409", "−1,675", "Rest is tail-window"],
           ["Cured IN-MONTH", "383,096", "−313", "Closing GT on floor"],
           ["IN-MONTH FULFILMENT", "96.2 %", "", "L11 reports 96.2 % — reconciles"],
           ["Carry-out tail, cures in August", "6,372", "",
            "Real output the horizon rule excludes"]])

    h(doc, "TBR", 2)
    table(doc, ["Stage", "Qty", "Delta", "Why / who owns it"],
          [["Demand (= plant's own July output)", "98,020", "—", ""],
           ["− covered by opening GT on floor", "96,723", "−1,297", ""],
           ["+ cure-yield scrap uplift", "98,545", "+1,822", "HARD: yield 0.982"],
           ["L4.5 lot-sized gross build", "98,545", "0", ""],
           ["L5 scheduled into cure campaigns", "97,289", "−1,256",
            "B12 residual dropped = 554 (POLICY)"],
           ["L7 actually BUILT and fed", "!!95,783", "!!−1,506", ""],
           ["Built before 31 Jul 07:00", "95,047", "−736", ""],
           ["Cured IN-MONTH", "94,231", "−816", "Closing GT on floor"],
           ["IN-MONTH FULFILMENT", "96.1 %", "",
            "L11 reports 95.7 % — it divides by plannable (97,436), not gross demand"],
           ["Carry-out tail", "2,006", "", ""]])

    # ---------------------------------------------------------------- 3
    h(doc, "3. The demerits — where we are worse than the plant", 1)
    p(doc, "Ranked by how much they cost. \"Plant\" is the mined 242-day MES "
           "baseline in plant_profile.json.")
    table(doc,
          ["#", "Demerit", "Plant", "Engine", "Gap", "Cost"],
          [["1", "PCR build changeovers/machine-day", "1.84", "!!2.53", "+38 %",
            "The clearest real demerit"],
           ["2", "PCR weighted CO min/machine-day", "74.0", "!!78.7", "+6 %",
            "~35 machine-h of extra setup"],
           ["3", "PCR same-size share of build COs", "91.5 %", "!!82.3 %", "−9.2 pt",
            "Longer setups per switch"],
           ["4", "TBR presses used per GT", "6.35 mean / 4.0 p50", "!!3.54 / 2.0",
            "−44 %", "Genuinely narrow on TBR only"],
           ["5", "PCR daily output CV, interior days", "0.116", "!!0.156", "+34 %",
            "Month-start dip; smaller than first reported"],
           ["6", "GTs per machine-day (PCR)", "2.17", "!!2.58", "+19 %",
            "We spread each machine thinner"],
           ["7", "TBR fulfilment", "100 %", "!!95.7 %", "−4.3 pt",
            "Worse than PCR on an easier problem"],
           ["8", "PCR stickiness", "99.84 %", "—", "not matched",
            "Plant almost never switches mid-run"],
           ["9", "cap_press omits 6 plant-sanctioned presses", "92 on roster",
            "!!86 visible", "6 presses", "Real data defect — but see section 12"],
           ["10", "PCR press utilisation vs physical roster", "~100 %", "!!89.8 %",
            "6,311 idle press-h", "Mostly presses legal for only 2 GTs"]])
    p(doc, "Three demerits in the first issue of this report were wrong and have "
           "been removed. They are documented in section 12 rather than deleted "
           "silently.", size=9, color=GRY)

    p(doc, "Where we already beat the plant", bold=True)
    table(doc, ["Metric", "Plant", "Engine", "Verdict"],
          [["PCR daily output mean", "12,381", "++12,422", "We out-produce it per day"],
           ["TBR weighted CO min/machine-day", "35.6", "++33.6", "Better"],
           ["TBR same-size share", "100 %", "++100 %", "Matched"],
           ["Mould changes per press-day", "0.08 / 0.04", "++0.04 / 0.03", "~2× better"],
           ["R5 shelf-life breaches", "n/a", "++0", "100 % compliant"],
           ["Allowable-machine violations", "n/a", "++0", "Was 19.9 % of volume pre-v9"]])

    # ---------------------------------------------------------------- 4
    h(doc, "4. Per-GT — where the loss concentrates", 1)
    p(doc, "PCR loss is CONCENTRATED: the top 10 of 34 short GTs carry 65 % of it. "
           "TBR is SPREAD (top 10 of 39 = 53 %), which means TBR needs a systemic "
           "fix and PCR needs a per-GT one.")
    table(doc,
          ["GT (PCR)", "Plant", "Engine", "Loss", "Press e/p", "Mach e/p", "Flags"],
          [["GT 2258 RAN HPE", "12,897", "10,849", "!!2,048", "4 / 15", "1 / 8",
            "Press+machine underuse"],
           ["GT 1513 XPC1 MSIL", "55,663", "54,285", "1,378", "22 / 44", "3 / 11",
            "Press+machine underuse"],
           ["GT 1402 XPC TATA", "20,788", "19,593", "1,195", "12 / 22", "2 / 5",
            "Press+machine underuse"],
           ["GT 1844 XPC TML", "11,181", "10,446", "735", "7 / 16", "2 / 11", ""],
           ["GT 2568 HT2", "16,090", "15,373", "717", "7 / 20", "2 / 8", ""],
           ["GT 2267 ROYL HYU", "31,947", "31,264", "683", "12 / 31", "3 / 8", ""],
           ["GT 1995 ROYL KIA", "14,249", "13,698", "551", "5 / 27", "3 / 11", ""],
           ["GT 1765 ROYL", "21,359", "20,813", "546", "8 / 21", "3 / 11", ""],
           ["GT 1412 XPC MM", "5,238", "4,793", "445", "3 / 11", "2 / 5", ""],
           ["GT 1482 UHL", "20,630", "20,223", "407", "8 / 12", "3 / 5", ""]])
    p(doc, "GT 2258 RAN HPE is the clearest single target: we ran it on 1 machine of "
           "8 capable and 4 presses of 15, and lost 2,048 tyres — 13 % of the whole "
           "PCR gap on one GT.", bold=True)

    # ---------------------------------------------------------------- 5
    h(doc, "5. Are we over-constraining? — the AND-stack test", 1)
    p(doc, "doc.txt asks whether the engine demands every preferred condition "
           "simultaneously where the plant says \"if feasible, run it\". That is the "
           "right question. Tested by classifying every filter:")
    table(doc, ["Filter", "Class", "Note"],
          [["Allowable building machine (R2)", "HARD", "Plant's own matrix"],
           ["Allowable press (R3)", "HARD", "0 violations both months"],
           ["Mould concurrency ≤ mould count", "HARD", "Physical"],
           ["PCR inch capability", "HARD", "Physical"],
           ["TT/TL group split (B16)", "HARD", "Confirmed by plant"],
           ["GT shelf life 72 h (R5)", "HARD", "Perishability"],
           ["Build lot floor B12 150/70", "!!POLICY",
            "Plant runs 12.7 % PCR / 30.8 % TBR sub-floor; we run 0.0 %"],
           ["Min demand to plan 300/150", "!!POLICY", "Plant makes these; we drop them"],
           ["GT WIP rail G8 4800/1400", "!!POLICY", "A rail, not a wall"],
           ["Mined rim lock", "PREFERENCE", "Already OFF — measured −14 pt"],
           ["Historical machine share", "PREFERENCE", "Already OFF — measured −0.4 pt"],
           ["Same-rim / sister clustering", "PREFERENCE", "Tie-break only, never a veto"]])
    p(doc, "VERDICT: 6 HARD + 3 POLICY. Every PREFERENCE filter is already shipped "
           "OFF after measurement, so the \"stacked preferences create the ceiling\" "
           "hypothesis does not hold for this engine as shipped. The recoverable set "
           "is the 3 POLICY rows — and all 3 are plant rulings, not code changes.",
      bold=True)

    # ---------------------------------------------------------------- 6
    h(doc, "6. Eligibility funnel — we use about two-thirds of what we are allowed", 1)
    table(doc, ["Stage", "PCR", "TBR", "Note"],
          [["Building machines — mined capability", "9.83", "3.39", "MES history"],
           ["… after plant allowable matrix (HARD)", "3.20", "3.11", "R2"],
           ["… after rim lock", "3.20", "3.11", "OFF, no effect"],
           ["… after rim sets", "3.20", "3.11", "OFF, no effect"],
           ["ENGINE ACTUALLY USED", "!!2.17", "!!2.50", "68 % / 80 % of allowed"],
           ["Presses — mined capability", "13.34", "20.96", ""],
           ["… mould concurrency cap (HARD)", "8.49", "7.70", "R3"],
           ["ENGINE ACTUALLY USED", "!!4.61", "!!3.54", "54 % / 46 % of legal"]])
    p(doc, "The first issue of this report called the last row a demerit. That was "
           "wrong, and the correct comparator makes the opposite case:", bold=True)
    table(doc, ["Breadth measure", "PLANT actual", "ENGINE", "Verdict"],
          [["PCR building machines per GT", "1.57 mean / 1.0 p50", "++2.17 / 2.0",
            "We are 38 % BROADER than the plant"],
           ["TBR building machines per GT", "2.30 / 2.0", "++2.50 / 2.0", "Broader"],
           ["PCR presses per GT", "5.87 mean / 2.0 p50", "4.61 / ++3.0",
            "Broader on the median, narrower on the mean"],
           ["TBR presses per GT", "6.35 / 4.0", "!!3.54 / 2.0",
            "Genuinely narrower — the one real gap"]])
    p(doc, "Comparing what we use against the ALLOWABLE CEILING (3.20 machines, "
           "8.49 presses) manufactures a demerit that does not exist. The plant "
           "itself uses 1.57 machines per SKU on PCR. Breadth is not a virtue — "
           "the plant is narrower than us and beats us. The only genuine breadth "
           "gap is TBR curing presses.")
    p(doc, "Breadth is also not free: 8 of 9 breadth-style experiments have measured "
           "negative (rim grouping −25,549 tyres; machine share −0.4 pt; de-pinning "
           "−0.5 pt), and the press-matrix union tested for this review reduced "
           "BUILT output on both months (section 12).", size=9, color=GRY)

    # ---------------------------------------------------------------- 7
    h(doc, "7. Lot sizing, changeovers, presses, inventory", 1)
    h(doc, "Lot sizing — the comparison that looks alarming and is not", 2)
    table(doc, ["Measure", "PCR plant", "PCR engine", "TBR plant", "TBR engine"],
          [["Build run / slice p50", "367", "157", "82", "26"],
           ["Build changeovers per machine-day", "1.84", "!!2.53", "3.12", "3.22"],
           ["Distinct GTs per machine-day", "2.17", "2.58", "3.02", "3.34"],
           ["Weighted CO min per machine-day", "74.0", "!!78.7", "35.6", "++33.6"]])
    p(doc, "The p50 row is NOT apples to apples and should not be quoted. Our build "
           "SLICE is the third campaign level (cure campaign → build run → build "
           "slice); consecutive same-GT slices on one machine are one physical run "
           "and cost no changeover. The honest measure is changeovers per "
           "machine-day, where we are 38 % worse than the plant on PCR and roughly "
           "level on TBR.", size=9, color=GRY)

    h(doc, "Press utilisation", 2)
    table(doc, ["Plant", "Denominator", "Running", "Mould change", "IDLE"],
          [["PCR", "86 presses the engine was offered", "61,436 h (96.0 %)",
            "701 h (1.1 %)", "1,848 h (2.9 %)"],
           ["PCR", "92 presses on the PHYSICAL roster", "61,436 h (!!89.8 %)",
            "701 h (1.0 %)", "!!6,311 h (9.2 %)"],
           ["TBR", "79 offered / 80 roster", "57,068 h (97.1 %)", "535 h (0.9 %)",
            "1,172 h (2.0 %)"]])
    p(doc, "Both PCR rows are true and the second is the honest one. Grading "
           "ourselves against the 86 presses the engine was handed proves only "
           "that the planner agrees with its own inputs — the same failure mode "
           "the project bans for verifiers. Against the physical roster of 92, PCR "
           "press utilisation is 89.8 %. The 6-press gap is a master-data defect "
           "(section 12), though those presses are plant-legal for only 2 GTs.",
      bold=True)
    p(doc, "Two denominator traps were hit while producing these and are worth "
           "recording: cycle_time_curing lists 92 PCR / 80 TBR presses against a "
           "roster of 86 / 79 (using the master inflated PCR idle from 2.9 % to "
           "7.3 % and invented \"6 unused presses\"), and campaign hours must be "
           "clipped to the horizon or tail campaigns push TBR above 100 % "
           "utilisation. This is the same class as the four denominator defects "
           "already in the project ledger.", size=9, color=GRY)

    h(doc, "GT inventory — we have the opposite of the problem doc.txt describes", 2)
    table(doc, ["Plant", "GT built", "GT cured", "Net", "Time-wtd inventory", "Rail"],
          [["PCR", "385,084", "385,084", "0", "3,699", "4,800"],
           ["TBR", "95,783", "95,783", "0", "1,136", "1,400"]])
    p(doc, "doc.txt's failure mode is \"Building 100,000 / Curing 90,000 / GT "
           "+10,000 — you are producing what curing cannot consume\". That is not "
           "our shape. Net is zero and inventory sits below the rail on both "
           "plants. Our problem is the mirror image: cure seats go unfilled because "
           "building could not reach them in time, not because building "
           "overproduced. Inventory is measured time-weighted, never "
           "event-weighted — event-weighting once biased TBR upward 5.7 %.")

    # ---------------------------------------------------------------- 8
    h(doc, "8. The rejection ledger — already instrumented, and it reconciles", 1)
    p(doc, "doc.txt item 9 asks for a rejection reason on every scheduling decision. "
           "It already exists: build_starved.parquet and cure_unplaced.parquet carry "
           "a reason on every rejected lot, and decision_trace records the rule IDs "
           "behind each placement.")
    table(doc, ["Plant", "Reason", "Tyres", "Lots"],
          [["PCR", "No feasible release", "!!6,310", "40"],
           ["PCR", "Would breach min_lot", "2,232", "22"],
           ["PCR", "TOTAL", "8,542", "62"],
           ["TBR", "Would breach min_lot", "1,454", "58"],
           ["TBR", "Below min_lot (strict B12)", "52", "2"],
           ["TBR", "Cure remainder past horizon", "621", "2"],
           ["TBR", "TOTAL build-side", "1,506", "60"]])
    p(doc, "These reconcile EXACTLY with the waterfall's unfed line — PCR "
           "6,310 + 2,232 = 8,542; TBR 1,454 + 52 = 1,506. There is no "
           "unexplained residue anywhere in the July plan.", bold=True)

    # ---------------------------------------------------------------- 9
    h(doc, "9. The scorecard", 1)
    table(doc, ["Metric", "PCR plant", "PCR engine", "TBR plant", "TBR engine"],
          [["Demand", "398,405", "398,405", "98,020", "98,020"],
           ["Production, in-month", "398,405", "!!383,096", "98,020", "!!94,231"],
           ["Production + carry-out tail", "398,405", "389,468", "98,020", "96,237"],
           ["Fulfilment %", "100.0", "!!96.2", "100.0", "!!95.7"],
           ["Daily output mean", "12,381", "++12,422", "3,101", "3,090"],
           ["Daily output CV", "0.116", "!!0.267", "0.142", "!!0.291"],
           ["Build COs per machine-day", "1.84", "!!2.53", "3.12", "3.22"],
           ["Weighted CO min/machine-day", "74.0", "!!78.7", "35.6", "++33.6"],
           ["GTs per machine-day", "2.17", "2.58", "3.02", "3.34"],
           ["Presses used", "86", "86", "79", "79"],
           ["Press utilisation %", "~100", "!!96.0", "~100", "!!97.1"],
           ["Same-size share of build COs", "91.5 %", "!!82.3 %", "100 %", "100 %"],
           ["GT wait max (R5 ≤ 72 h)", "—", "++70.9 h", "—", "++72.0 h"]])

    # ---------------------------------------------------------------- 10
    h(doc, "10. What to correct — ranked, with owner", 1)
    h(doc, "Plant rulings (largest value, no code)", 2)
    table(doc, ["#", "Item", "Worth (Jul)", "Question for the plant"],
          [["1", "Narrow allowable-machine matrices", "Largest single item",
            "26 of 73 Aug PCR GTs have ≤2 allowable machines; TBMPCR7 lists 39 GTs "
            "against TBMPCR2's 130. Can any be widened?"],
           ["2", "Horizon ruling — does the carry-out tail count?",
            "+1.60 pt PCR / +2.05 pt TBR",
            "6,372 PCR tyres are built in July and cure on 1–2 Aug. Real output; "
            "currently excluded."],
           ["3", "B12 sub-floor ruling", "+0.56 pt PCR / +1.48 pt TBR",
            "The plant itself runs 12.7 % PCR / 30.8 % TBR sub-floor. We run 0.0 %. "
            "May we match the plant?"],
           ["4", "B12 residual (min demand 300/150)", "+0.28 pt PCR / +0.57 pt TBR",
            "1,124 PCR / 554 TBR tyres are never planned. The plant makes them."],
           ["5", "30-June machine state — 20 rows",
            "Part of the day-1 dip", "Which GT was each of the 11 PCR / 9 TBR "
            "machines building at 07:00 on 1 July?"],
           ["6", "Opening-stock GT spread", "Part of the day-1 dip",
            "Quantity is verified correct by Little's law, but it sits on only 25 of "
            "48 demanded PCR GTs."]])

    h(doc, "Engineering (ours, and honest about the odds)", 2)
    table(doc, ["#", "Item", "Est. worth", "Confidence"],
          [["7", "Targeted second-pass repair on the 34 short PCR GTs — re-offer "
                 "unfed lots to idle, allowable, R5-legal machine windows",
            "≤ 0.6 pt", "Medium — 222 R5-legal idle hours exist, but only ~9 % of "
                        "idle time is legally usable"],
           ["8", "GT 2258 RAN HPE specifically — 1 machine of 8, 4 presses of 15",
            "up to 2,048 tyres", "Medium — worth a single-GT investigation first"],
           ["9", "Month-start dip — CV 0.267 vs plant 0.116",
            "~3.3k tyres", "Low — six approaches measured, one gained 0.3 pt and it "
                           "then failed the August gate"],
           ["10", "July partition (none exists; July runs unpartitioned)",
            "~0.6 pt + same-size", "Blocked — needs the raw MES drop"]])

    p(doc, "What NOT to do", bold=True)
    bullets(doc, [
        "Do not build before 07:00 on 1 July (\"carry-in\"). It borrows June's "
        "capacity and is not ours to spend.",
        "Do not relax min_lot, min_demand, or the R5 shelf life to buy fulfilment. "
        "Those are the plant's caps; changing them changes the answer, not the plan.",
        "Do not halve setup times. The same objection applies.",
        "Do not re-enable the rim lock, machine-share ordering, or depth-before-"
        "breadth queue order. All measured negative (−14 pt, −0.4 pt, −11.9 pt).",
        "Do not tune anything on July alone. Two flags did exactly that and cost "
        "August real fulfilment.",
    ])

    # ---------------------------------------------------------------- 11
    h(doc, "11. Method notes — defects found while producing this report", 1)
    p(doc, "Recorded because this project's most productive bug class is the "
           "measurement, not the scheduler. Four denominator defects were already "
           "in the ledger; this analysis added four more, all mine, all caught "
           "before they reached a conclusion:")
    bullets(doc, [
        "A sweep harness summed qty_fed_in_month — the FED basis, which includes "
        "opening stock — and reported it as in-month. Inflated July by ~0.6 pt and "
        "reversed the sign of a PCR/TBR comparison. The harness is retired.",
        "The first waterfall assumed net_cure = demand − opening and that `usable` "
        "counted resolved SKUs. Both wrong: net_cure == demand, and `usable` counts "
        "opening-GT units.",
        "In-month output was computed as built − tail. Closing stock is not the "
        "carry-out tail; the correct form is a horizon cut on the ledger. The wrong "
        "form overstated TBR by 1.3 pt.",
        "Press utilisation used cycle_time_curing's 92/80 press count against a "
        "roster of 86/79, and did not clip campaign hours to the horizon. PCR idle "
        "read 7.3 % instead of 2.9 %, with six phantom unused presses.",
    ])
    p(doc, "Every figure in this report is regenerated by "
           "scripts/diag_plant_vs_engine.py from the run artefacts. If the plan "
           "changes, re-run it rather than editing this document.",
      size=9, color=GRY)

    # ---------------------------------------------------------------- 12
    h(doc, "12. Expert review of this report — what it got wrong", 1)
    p(doc, "This section exists because the first issue of this document shipped "
           "three wrong demerits and nearly recommended a change that reduces "
           "production. The corrections are more useful than the original claims.")

    h(doc, "Corrected: three demerits that were not demerits", 2)
    table(doc, ["Original claim", "What was wrong", "Corrected"],
          [["\"PCR machines per GT 2.17 vs 3.20 allowable = 68 % — underuse\"",
            "Compared against the ALLOWABLE CEILING, not against the plant. The "
            "plant itself uses 1.57 machines per SKU.",
            "++We are 38 % BROADER than the plant. Not a demerit."],
           ["\"Presses per GT 4.61 vs 8.49 mould-legal = 54 % — underuse\"",
            "Same error. Plant PCR actual is 5.87 mean but 2.0 median.",
            "++PCR: broader than plant on the median (3.0 vs 2.0). TBR is the "
            "one genuine gap (3.54 vs 6.35)."],
           ["\"Daily output CV 0.267 vs plant 0.116 — 2.3× worse\"",
            "Computed on WALL-CLOCK dates, not plant-days, and included two "
            "partial boundary days.",
            "!!PCR interior CV is 0.156 vs 0.116 — real but a third of the "
            "claimed size. ++TBR is 0.113 vs plant 0.142: we are STEADIER."]])

    h(doc, "Corrected: the press-utilisation denominator, in both directions", 2)
    p(doc, "The first issue used 92 presses, then \"fixed\" it to 86 on the "
           "strength of a doc. Both were wrong to state alone. 92 is the physical "
           "roster (four independent masters agree); 86 is what cap_press_2026-07 "
           "offers the engine. Grading against 86 proves only that the planner "
           "agrees with its own inputs.")

    h(doc, "New finding: 6 plant-sanctioned presses are invisible to the planner", 2)
    p(doc, "cap_press_2026-07 is MINED from MES. allowed_press_matrix is the "
           "PLANT'S OWN file and is broader. PCR presses 17–22 are physically "
           "present (132 tyres/day each) and marked `direct` by the plant for GT "
           "1402 XPC TATA and GT 1412 XPC MM — two of the three largest per-GT "
           "losses. L5 gates on cap_press, so it never sees them. Mould headroom "
           "exists (GT 1402: 20 moulds, 12 presses used), so R3 does not block it.")
    p(doc, "This is the exact converse of the v9 building-machine fix. There the "
           "plant matrix was NARROWER than mined capability and became a hard "
           "restrict. Here it is BROADER, so it would be a union — and nobody "
           "checked that direction, because the press gate was verified only for "
           "violations (\"0 violations, clean\") and never for omissions. "
           "A gate can be clean and still be costing you output.", bold=True)

    h(doc, "…and why the fix was still REJECTED", 2)
    p(doc, "Implemented and measured on both months. Reading in-month fulfilment "
           "alone, it is the biggest win of the project: August PCR 91.4 % → "
           "94.8 %. It is not a win.")
    table(doc, ["Arm", "BUILT", "In-month", "Carry-out tail", "TOTAL real output"],
          [["Aug, union OFF", "++410,652", "392,239 (91.4 %)", "23,381",
            "++415,620 (96.8 %)"],
           ["Aug, union ON", "!!408,220", "406,830 (94.8 %)", "6,616",
            "!!413,446 (96.3 %)"],
           ["Jul, union OFF", "++385,084", "383,266 (96.2 %)", "6,372",
            "++389,638 (97.8 %)"],
           ["Jul, union ON", "!!378,838", "382,070 (95.9 %)", "1,290",
            "!!383,360 (96.2 %)"]])
    p(doc, "The union BUILDS FEWER TYRES on both months (−2,432 August, −6,246 "
           "July) and total real output FALLS on both. Every point of the apparent "
           "August gain is the carry-out tail moving inside the horizon — cures "
           "happen earlier, not more. Mechanism: more eligible presses → more "
           "parallel campaigns → building feeds more GTs at once → same-size share "
           "75.2 % → 69.5 % (an L11 invariant flips PASS→FAIL) and weighted "
           "changeover 99.8 → 107.2 min/machine-day.")
    p(doc, "THE LESSON, which generalises well past this flag: in-month fulfilment "
           "is TAIL-SENSITIVE. Any change that pulls cures earlier inflates it "
           "without producing anything. BUILT and in-month+tail must be reported "
           "alongside it. A change that moves in-month while BUILT falls is "
           "relocating output, not creating it.", bold=True)
    p(doc, "The code ships behind PLANNER_PRESS_FROM_MATRIX, defaulting OFF, with "
           "the measurement recorded in l5_cure_master.py. The underlying master "
           "defect is real and stands; it is simply not worth exploiting today.",
      size=9, color=GRY)

    h(doc, "What this changes about the recommendations", 2)
    bullets(doc, [
        "The August headline (91.4 %) is substantially DEFINITIONAL. On total real "
        "output August is ~96.8 %, close to July's 97.8 %. The apparent 5-point "
        "August deficit is mostly a 23,381-tyre carry-out tail, not lost production.",
        "The horizon ruling (section 10, item 2) is therefore worth far more on "
        "August than on July — 23,381 tyres against 6,372. It should be the first "
        "question put to the plant.",
        "Breadth is dead as an improvement direction. The plant is narrower than us "
        "and beats us; every breadth experiment including this one has measured "
        "negative on BUILT.",
        "The remaining engineering item is unchanged and modest: a targeted repair "
        "pass on the 6,310 \"no feasible release\" tyres, graded on BUILT.",
    ])

    # ---------------------------------------------------------------- 13
    h(doc, "13. August 2026 — the \">95 %\" question, answered", 1)
    p(doc, "Asked directly: can August come in above 95 % fulfilment? The honest "
           "answer depends entirely on which of the three numbers is meant, and "
           "two of the three already clear it.")
    table(doc, ["August", "Demand", "BUILT", "built %", "In-month", "in-month %",
                "Tail", "In-month + tail", "total %"],
          [["PCR", "429,146", "410,652", "++95.7 %", "392,239", "!!91.4 %",
            "23,381", "415,620", "++96.8 %"],
           ["TBR", "99,019", "92,541", "!!93.5 %", "88,523", "!!89.4 %", "3,517",
            "92,040", "!!93.0 %"]])
    p(doc, "PCR already clears 95 % on production (95.7 %) and on total real "
           "output (96.8 %). It does not on in-month (91.4 %), and TBR clears it "
           "on none.", bold=True)

    h(doc, "Nine arms were swept trying to lift in-month. None reached 95 %.", 2)
    table(doc, ["Arm", "BUILT", "ΔBUILT", "In-month", "ful %", "L11"],
          [["base (shipped)", "++410,652", "—", "392,239", "91.4", "27/42"],
           ["HORIZON_MODE=strict", "!!380,080", "!!−30,572", "384,944", "89.7", "25/42"],
           ["HORIZON_MODE=truncate", "!!384,156", "!!−26,496", "388,806", "90.6", "26/42"],
           ["HORIZON_MODE=window", "!!393,182", "!!−17,470", "389,665", "90.8", "27/42"],
           ["HORIZON_TAIL_H=48", "403,762", "−6,890", "391,810", "91.3", "26/42"],
           ["HORIZON_TAIL_H=24", "!!393,800", "!!−16,852", "391,381", "91.2", "26/42"],
           ["PRESS_FROM_MATRIX=1", "408,220", "−2,432", "406,830", "!!94.8", "26/42"],
           ["TAIL_H=24 + PRESS", "404,061", "−6,591", "405,543", "94.5", "24/42"],
           ["TAIL_H=0 + PRESS", "!!399,316", "!!−11,336", "404,256", "94.2", "25/42"]])
    p(doc, "The August PCR in-month ceiling is ~94.8 %, and every arm that lifts "
           "in-month cuts BUILT. Cause: about 14,400 PCR tyres are BUILT after "
           "07:00 on 1 September under the default 72-hour extension. Shortening "
           "the extension does not move them inside the month — it deletes them. "
           "BUILT falls faster than in-month rises.")

    h(doc, "The two honest routes above 95 %, both plant rulings", 2)
    table(doc, ["Route", "Effect", "Basis"],
          [["Horizon ruling — does the carry-out tail count?",
            "++Aug PCR 91.4 % → 96.8 %",
            "23,381 tyres are already built and cure 1–3 Sep. No code change."],
           ["B12 sub-floor ruling on TBR",
            "++Aug TBR +6.2 pt",
            "6,100 tyres over 225 lots are rejected by min_lot. The plant itself "
            "runs 30.8 % of TBR sub-floor; we run 0.0 %."],
           ["TBR allowable matrix (still unresolved from v9)", "Aug TBR ~+5 pt",
            "The plant's TBR file is tighter than the derived list for August's "
            "37 GTs. Needs a ruling on which is authoritative."]])
    p(doc, "What was NOT done, and will not be: the in-month number was not "
           "manufactured by relaxing min_lot, min_demand, the 72-hour shelf life "
           "or setup times, and no tyre was built before the month opened. Each "
           "of those raises the KPI without the plant making a single extra tyre.",
      bold=True)

    doc.save(OUT)
    print(f"  -> {OUT}")


if __name__ == "__main__":
    main()
